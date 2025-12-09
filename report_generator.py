import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import re

def create_word_report(title, strategy_name, content, metrics):
    """
    Create a Word document report.
    """
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Strategy Name
    doc.add_heading(f"Strategy: {strategy_name}", level=1)
    
    # Metrics Section
    doc.add_heading("Key Metrics", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    for key, value in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key.replace('_', ' ').title()
        if isinstance(value, float):
            row_cells[1].text = f"{value:.2f}"
        else:
            row_cells[1].text = str(value)
            
    # Analysis Content
    doc.add_heading("AI Analysis", level=2)
    
    # Simple Markdown parsing for the content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('**') and line.endswith('**'):
            # Bold heading
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        elif line.startswith('1. '):
             doc.add_paragraph(line, style='List Number')
        else:
            # Regular text, check for bold parts
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.bold = True
                else:
                    p.add_run(part)

    # Save to IO stream
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'ZMtech AI Options Calculator Pro - Report', 0, 1, 'C')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def sanitize_text(text):
    """
    Replace incompatible unicode characters with ASCII equivalents for FPDF.
    """
    replacements = {
        '\u2013': '-',  # En dash
        '\u2014': '--', # Em dash
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2022': '*',  # Bullet
        '\u2026': '...', # Ellipsis
        '\u00a0': ' ',  # Non-breaking space
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
        
    # Remove any other non-latin-1 characters
    return text.encode('latin-1', 'replace').decode('latin-1')

def create_pdf_report(title, strategy_name, content, metrics):
    """
    Create a PDF report.
    """
    pdf = PDF()
    pdf.add_page()
    
    # Strategy Name
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, f"Strategy: {sanitize_text(strategy_name)}", 0, 1)
    pdf.ln(5)
    
    # Metrics
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, "Key Metrics", 0, 1)
    pdf.set_font('Arial', '', 10)
    
    for key, value in metrics.items():
        key_str = key.replace('_', ' ').title()
        if isinstance(value, float):
            val_str = f"{value:.2f}"
        else:
            val_str = str(value)
        pdf.cell(0, 6, sanitize_text(f"{key_str}: {val_str}"), 0, 1)
    
    pdf.ln(10)
    
    # Analysis
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, "AI Analysis", 0, 1)
    pdf.set_font('Arial', '', 10)
    
    # Simple Markdown parsing
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            pdf.ln(2)
            continue
            
        # Sanitize line
        line = sanitize_text(line)
            
        # Handle headers
        if line.startswith('#'):
            clean_line = line.lstrip('#').strip()
            pdf.set_font('Arial', 'B', 11)
            pdf.multi_cell(0, 6, clean_line)
            pdf.set_font('Arial', '', 10)
        elif line.startswith('**') and line.endswith('**'):
             clean_line = line.replace('**', '')
             pdf.set_font('Arial', 'B', 10)
             pdf.multi_cell(0, 6, clean_line)
             pdf.set_font('Arial', '', 10)
        else:
            # Remove bold markers for regular text
            clean_line = line.replace('**', '')
            pdf.multi_cell(0, 5, clean_line)
            
    # Output to IO stream
    try:
        return pdf.output(dest='S').encode('latin-1')
    except:
        # Fallback for newer FPDF versions or if encoding fails
        # If dest='S' returns a string, we encode it. 
        # If it returns bytes (some versions), we return as is.
        out = pdf.output(dest='S')
        if isinstance(out, str):
            return out.encode('latin-1', 'replace')
        return out 

